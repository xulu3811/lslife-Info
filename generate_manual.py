import os
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT, WD_LINE_SPACING
from docx.oxml import OxmlElement, ns

def add_page_number(run):
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(ns.qn('w:fldCharType'), 'begin')
    instrText = OxmlElement('w:instrText')
    instrText.set(ns.qn('xml:space'), 'preserve')
    instrText.text = 'PAGE'
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(ns.qn('w:fldCharType'), 'end')
    run._r.append(fldChar1)
    run._r.append(instrText)
    run._r.append(fldChar2)

def set_font(run, font_name):
    run.font.name = font_name
    r = run._element
    rFonts = r.find(ns.qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        r.append(rFonts)
    rFonts.set(ns.qn('w:eastAsia'), font_name)

doc = Document()
# set margins
for section in doc.sections:
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1.2)
    section.right_margin = Inches(1.2)
    
    header = section.header
    header_para = header.paragraphs[0]
    header_para.paragraph_format.tab_stops.add_tab_stop(Inches(6.0))
    
    run_left = header_para.add_run('同城•连山 V1.0\t')
    set_font(run_left, 'Microsoft YaHei')
    
    run_page = header_para.add_run()
    add_page_number(run_page)
    set_font(run_page, 'Microsoft YaHei')

def add_heading(text, level=1):
    p = doc.add_heading(text, level=level)
    for run in p.runs:
        set_font(run, 'Microsoft YaHei')
        run.font.color.rgb = RGBColor(0, 0, 0)
    return p

def add_paragraph(text, bold=False):
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 1.5
    run = p.add_run(text)
    run.bold = bold
    set_font(run, 'SimSun')
    run.font.size = Pt(12)
    return p

def add_screenshot_placeholder(desc):
    p = doc.add_paragraph()
    p.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(12)
    run = p.add_run(f'【在此处插入截图：{desc}】')
    run.bold = True
    run.font.color.rgb = RGBColor(255, 0, 0)
    set_font(run, 'Microsoft YaHei')
    run.font.size = Pt(14)
    
    # Add some empty space to simulate the image height
    for _ in range(5):
        add_paragraph('')

# Cover
doc.add_paragraph('\n\n\n\n\n')
title = doc.add_paragraph('同城•连山 V1.0')
title.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
run = title.runs[0]
run.bold = True
run.font.size = Pt(36)
set_font(run, 'Microsoft YaHei')

subtitle = doc.add_paragraph('用户操作手册')
subtitle.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
run = subtitle.runs[0]
run.bold = True
run.font.size = Pt(24)
set_font(run, 'Microsoft YaHei')
doc.add_page_break()


# 1. 简介
add_heading('第一章 软件简介与环境要求', 1)
add_paragraph('“同城•连山” 是一款专为中国县域级下沉市场量身打造的本地同城分类信息与生活服务平台。本系统旨在打破下沉市场的信息壁垒，为本地居民提供二手闲置、房源租售、求职招聘等一站式本地生活信息服务。')
add_paragraph('1.1 运行环境要求', bold=True)
add_paragraph('客户端：Android 7.0 (API 24) 及以上版本的智能手机。')
add_paragraph('服务端：Ubuntu 24.04 操作系统，支持 Node.js 与 PostgreSQL 数据库。')
doc.add_page_break()

# 2. 登录注册
add_heading('第二章 用户注册与登录', 1)
add_paragraph('本章节主要介绍用户如何通过手机终端进入系统，并完成安全认证。')
add_paragraph('2.1 手机号验证码登录/注册', bold=True)
add_paragraph('用户首次打开 APP 时，将进入欢迎页及登录界面。输入正确的 11 位手机号码后，点击获取验证码。系统将通过短信通道向用户下发 6 位动态验证码。')
add_paragraph('用户输入验证码并同意《用户服务协议》及《隐私政策》后，点击登录。若为首次登录的手机号，系统将在后台自动为其创建账号并分配默认头像与昵称。')
add_screenshot_placeholder('APP登录注册界面 (输入手机号及验证码)')
add_paragraph('2.2 资料完善', bold=True)
add_paragraph('登录成功后，用户可前往“我的”页面修改个人资料，包括自定义头像、修改高辨识度昵称等，相关资料修改需经平台后台人工审核通过后方可全网展示。')
doc.add_page_break()

# 3. 信息瀑布流浏览
add_heading('第三章 首页与信息瀑布流浏览', 1)
add_paragraph('首页是平台的核心流量分发阵地，采用现代化的 3D Soft UI 极简美学设计及双列瀑布流排版引擎。')
add_paragraph('3.1 首页金刚区与分类导航', bold=True)
add_paragraph('首页顶部设有多功能搜索栏与轮播 Banner。中部的“金刚区”提供房产、招聘、二手、服务等核心频道的快捷入口。用户点击任意图标即可进入对应的二级分类列表域。')
add_screenshot_placeholder('首页金刚区与Banner展示区')
add_paragraph('3.2 双列信息瀑布流浏览', bold=True)
add_paragraph('在首页下半部分及各大二级分类中，平台采用小红书式的双列瀑布流卡片展示信息。卡片直观呈现物品/服务主图、高亮标题、动态属性（如“99新”、“急售”）及直观的价格标识。')
add_paragraph('用户可以通过上下滑动进行沉浸式浏览，系统会自动通过后台 API 进行分页加载，保证浏览的丝滑度与极低的网络延迟。')
add_screenshot_placeholder('首页或同城发现页的双列瀑布流商品卡片')
add_paragraph('3.3 详情页展示', bold=True)
add_paragraph('点击瀑布流中的任意卡片即可进入信息详情页。详情页支持高清组图滑动浏览，完整展示发布者的地理位置、具体参数及详细描述。页面底部提供“收藏”、“电话联系”与“在线私聊”功能入口。')
add_screenshot_placeholder('信息发布详情页 (展示图片、价格与聊天入口)')
doc.add_page_break()

# 4. 动态表单发布
add_heading('第四章 动态表单与信息发布', 1)
add_paragraph('为了适配多样化的本地生活需求，系统内置了高度可定制的动态属性表单（Dynamic Schema）。')
add_paragraph('4.1 选择发布类目', bold=True)
add_paragraph('用户点击底部导航栏中间的“发布”按钮（大加号），即可呼出类目选择面板。用户需根据所发布的物品或服务性质，精确选择对应的一级与二级分类。')
add_screenshot_placeholder('发布类目选择面板')
add_paragraph('4.2 填写动态属性表单', bold=True)
add_paragraph('系统将根据用户选择的分类，实时渲染出不同的输入项。例如：选择“二手手机”，系统要求必填“品牌”、“成色”、“内存容量”；若选择“房屋出租”，则表单自动切换为“户型”、“面积”、“朝向”、“月租金”等核心字段。')
add_paragraph('用户还可以批量上传多张高清实物图片，系统在本地完成图片无损压缩后进行极速上传。填写完毕并确认信息无误后，点击“立即发布”。')
add_screenshot_placeholder('填写发布信息的动态表单界面')
doc.add_page_break()

# 5. IM即时通讯
add_heading('第五章 即时通讯与消息互动 (IM)', 1)
add_paragraph('本平台内置了基于 WebSocket 长连接机制的全双工即时通讯系统，支持买卖双方或服务供需双方的高频互动。')
add_paragraph('5.1 消息列表页', bold=True)
add_paragraph('在“消息”标签页中，用户可以查看所有的历史会话列表，包括系统通知、交易提醒以及与其他用户的私信记录。系统会在头像右上角呈现醒目的未读消息红点徽章。')
add_screenshot_placeholder('消息列表页面 (展示会话与未读红点)')
add_paragraph('5.2 一对一即时通讯聊天页', bold=True)
add_paragraph('在聊天室内，用户可以发送文字、图片等多媒体消息。为促进交易流转，当用户从某条发布信息详情页发起咨询时，系统会自动向对方发送一张“信息快照卡片”，便于双方明确沟通主体。所有聊天记录均经过 AES 加密及防篡改哈希上链，确保沟通安全。')
add_screenshot_placeholder('IM 即时通讯聊天页 (展示文字或商品快照卡片)')
doc.add_page_break()

# 6. 后台Dashboard
add_heading('第六章 平台运营与后台 Dashboard 管理', 1)
add_paragraph('为了构建绿色、健康、合规的网络空间，“同城•连山”配备了严格的“先审后发”风控引擎与移动端统一管理员工作流。')
add_paragraph('6.1 我的页面与运营入口', bold=True)
add_paragraph('平台管理员或超级运营账号在登录后，其“我的”页面功能区将开放专属的“平台运营与管理”板块。点击“待审批事项”即可进入管理中枢。')
add_screenshot_placeholder('“我的”页面 (包含平台运营与管理入口)')
add_paragraph('6.2 统一 Dashboard 审批面板', bold=True)
add_paragraph('系统将管理事务高度聚合，Dashboard 面板清晰展示了四大数据维度：待审帖子、待审个人资料、实名认证审核、商家入驻审核。各模块均配有直观的待办数量统计。')
add_screenshot_placeholder('后台 Dashboard 聚合审批面板')
add_paragraph('6.3 商家与实名认证审批流', bold=True)
add_paragraph('以“商家入驻审批”为例，管理员点击进入列表后，可查阅用户提交的营业执照、门头照片及法人信息。经过严格核实后，管理员点击“通过”，系统后台将自动为其静默建立店铺并赋予商家蓝V标识，完成商业流转自动化闭环。')
add_screenshot_placeholder('商家入驻或个人实名认证的详细审核页')

doc_path = r'C:\Users\xl246\.gemini\antigravity\brain\89736c11-396e-4c8c-b7da-e4d2ad973d64\LsLife_User_Manual.docx'
doc.save(doc_path)
print('DOCX generated successfully at:', doc_path)
