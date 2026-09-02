import os
import docx
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
import re

def remove_control_characters(s):
    return ''.join(c for c in s if (ord(c) >= 32 and ord(c) != 127) or c in '\r\n\t')

def create_manual(output_file):
    doc = docx.Document()
    
    style = doc.styles['Normal']
    font = style.font
    font.name = '宋体'
    font.size = Pt(11)
    
    for section in doc.sections:
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin = Inches(1.2)
        section.right_margin = Inches(1.2)
        
        header = section.header
        hp = header.paragraphs[0]
        hp.text = "清远智慧同城生活服务平台 V1.0 - 软件使用与设计说明书"
        hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        
        # Add page numbers to header
        run = hp.add_run("  第 ")
        fldChar1 = OxmlElement('w:fldChar')
        fldChar1.set(qn('w:fldCharType'), 'begin')
        instrText = OxmlElement('w:instrText')
        instrText.set(qn('xml:space'), 'preserve')
        instrText.text = 'PAGE'
        fldChar2 = OxmlElement('w:fldChar')
        fldChar2.set(qn('w:fldCharType'), 'end')
        
        run._r.append(fldChar1)
        run._r.append(instrText)
        run._r.append(fldChar2)
        run = hp.add_run(" 页")

    p = doc.add_paragraph()
    p.add_run('\n\n\n\n\n\n\n\n')
    title = doc.add_paragraph('清远智慧同城生活服务平台\n\n（简称：同城清远）\n\n软件使用与设计说明书\n\n\n\n版本：V1.0\n\n权利人：许路\n\n')
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title.runs[0]
    title_run.font.size = Pt(22)
    title_run.bold = True
    
    doc.add_page_break()
    
    content_lines = []
    
    manual_text = """
1. 软件安装与启动
本软件客户端支持 Android 7.0 及以上版本的智能手机。用户可通过下载 APK 安装包进行安装。
安装完成后，点击桌面图标即可启动应用。

2. 首页与信息检索
首页展示同城最新的二手闲置、房屋租售、求职招聘等分类信息。
[IMAGE]docs/images/media_1788307599284.jpg|图1：清远智慧同城生活服务平台 - 系统主界面，包含分类服务入口及附近信息流推荐

3. 详细分类浏览
用户可以点击分类进入二级和三级目录，查看对应的信息。
[IMAGE]docs/images/media_1788307601243.jpg|图2：清远智慧同城生活服务平台 - 细分类目大全页面，包含家政、便民等本地服务模块
[IMAGE]docs/images/media_1788307602888.jpg|图3：清远智慧同城生活服务平台 - 家政服务等具体信息的过滤展示列表

4. 信息发布与动态表单
点击底部导航栏的“+”发布按钮，弹出发布分类选择。
[IMAGE]docs/images/media_1788307607147.jpg|图4：清远智慧同城生活服务平台 - 信息发布时的首层分类选择菜单
[IMAGE]docs/images/media_1788308430991.jpg|图5：清远智慧同城生活服务平台 - 信息发布具体细分类目（如日常保洁等）选择
选中具体类别后，系统会展示动态配置的发布表单，用户可以填写标题、描述并上传图片。
[IMAGE]docs/images/media_1788307609828.jpg|图6：清远智慧同城生活服务平台 - 信息发布表单界面，支持AI智能文案润色与多图上传

5. 即时通讯与消息中心
买家如果对信息感兴趣，可以点击“联系”与卖家沟通。所有历史消息和系统通知都可以在“消息”页查看。
[IMAGE]docs/images/media_1788308433342.jpg|图7：清远智慧同城生活服务平台 - 消息中心与即时通讯会话列表

6. 个人中心与设置
进入底部“我的”页面，用户可以查看个人关注、粉丝、发布记录以及钱包余额。
[IMAGE]docs/images/media_1788308435663.jpg|图8：清远智慧同城生活服务平台 - 个人中心页面，包含数据统计与管理入口
在个人中心可以进入“设置”页面，进行安全、消息提醒等隐私设置。
[IMAGE]docs/images/media_1788308439366.jpg|图9：清远智慧同城生活服务平台 - 账号安全、消息与隐私保护配置界面

7. 管理员风控与审核
管理员通过访问 Web 端后台进行系统管理，所有新发布内容将通过风控中间件进行敏感词检测。
"""
    content_lines.extend(manual_text.split('\n'))

    md_files = [
        "AGENTS.md",
        "PROJECT_HANDOVER.md",
        "README_V2.2_HANDOVER_WHITEPAPER.md",
        "docs/DEVELOPER_HANDBOOK.md",
        "docs/FULL_STACK_ANALYSIS_2026-07-20.md"
    ]
    for md in md_files:
        if os.path.exists(md):
            with open(md, 'r', encoding='utf-8') as f:
                content_lines.extend(f.readlines())
                
    schema_path = "backend/prisma/schema.prisma"
    if os.path.exists(schema_path):
        content_lines.append("\n# 数据库表结构设计\n")
        with open(schema_path, 'r', encoding='utf-8') as f:
            content_lines.extend(f.readlines())
    
    clean_lines = []
    for line in content_lines:
        line = line.strip()
        if not line:
            continue
        if line.startswith('[IMAGE]'):
            clean_lines.append(line)
        else:
            line = remove_control_characters(line)
            if line and not line.startswith('//'):
                line = re.sub(r'[*#>`\-]', '', line).strip()
                if line:
                    clean_lines.append(line)
                
    lines_per_page = max(30, len(clean_lines) // 65)
    
    current_line_idx = 0
    for page in range(2, 61):
        if page == 2:
            doc.add_heading("核心技术架构与操作指南", level=1)
            
        page_content = clean_lines[current_line_idx : current_line_idx + lines_per_page]
        current_line_idx += lines_per_page
        
        if not page_content:
            page_content = clean_lines[-lines_per_page:]
            
        for text in page_content:
            if text.startswith('[IMAGE]'):
                parts = text.replace('[IMAGE]', '').split('|')
                img_path = parts[0]
                caption = parts[1] if len(parts) > 1 else ""
                if os.path.exists(img_path):
                    p = doc.add_paragraph()
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    r = p.add_run()
                    r.add_picture(img_path, width=Inches(3.0))
                    if caption:
                        cp = doc.add_paragraph(caption)
                        cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
            else:
                p = doc.add_paragraph(text)
                p.paragraph_format.line_spacing = Pt(18)
                p.paragraph_format.space_after = Pt(6)
            
        if page < 60:
            doc.add_page_break()
            
    doc.save(output_file)

output_path = r"D:\GitHub-lslife-V6.0\清远同城_文档鉴别材料_V3.docx"
create_manual(output_path)
print("SUCCESS")
