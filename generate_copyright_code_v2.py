import os
import docx
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

def collect_code(target_dir, ext_list):
    lines = []
    for root, dirs, files in os.walk(target_dir):
        for file in files:
            if any(file.endswith(ext) for ext in ext_list):
                filepath = os.path.join(root, file)
                try:
                    with open(filepath, 'r', encoding='utf-8') as f:
                        for line in f:
                            line = line.rstrip()
                            stripped = line.strip()
                            # skip empty and pure comment lines
                            if not stripped or stripped.startswith('//') or stripped.startswith('/*') or stripped.startswith('*'):
                                continue
                            if len(line) > 80:
                                line = line[:80] + "..."
                            lines.append(line)
                except Exception:
                    pass
    return lines

def create_docx(first_half, second_half, output_file):
    doc = docx.Document()
    
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(0.8)
        section.bottom_margin = Inches(0.8)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)
        
        header = section.header
        header_para = header.paragraphs[0]
        header_para.text = "清远智慧同城生活服务平台 V1.0"
        header_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        
        run = header_para.add_run("  第 ")
        fldChar1 = OxmlElement('w:fldChar')
        fldChar1.set(qn('w:fldCharType'), 'begin')
        instrText = OxmlElement('w:instrText')
        instrText.set(qn('xml:space'), 'preserve')
        instrText.text = 'PAGE'
        fldChar2 = OxmlElement('w:fldChar')
        fldChar2.set(qn('w:fldCharType'), 'end')
        
        run._r.append(fldChar1)
        run._r.append(instrText)
        run._r.append(fldChar2)
        run = header_para.add_run(" 页")

    # take exact 1500 lines for the first 30 pages and 1500 lines for the last 30 pages
    first_1500 = first_half[:1500]
    last_1500 = second_half[-1500:]
    
    final_lines = first_1500 + last_1500
    
    if len(final_lines) < 3000:
        print(f"Error: only collected {len(final_lines)} lines, need 3000.")
        return
        
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Courier New'
    font.size = Pt(9)
    rFonts = font.element.rPr.rFonts
    if rFonts is not None:
        rFonts.set(qn('w:eastAsia'), '宋体')

    for page in range(60):
        page_lines = final_lines[page*50 : (page+1)*50]
        
        p = doc.add_paragraph()
        p.paragraph_format.line_spacing = Pt(12) 
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.space_before = Pt(0)
        
        p.add_run('\n'.join(page_lines))
        
        if page < 59:
            doc.add_page_break()
            
    doc.save(output_file)

# 1. Collect Android Client code (First 30 pages)
android_lines = collect_code("android/app/src/main/java", [".kt"])
# 2. Collect Backend Server code (Last 30 pages)
backend_lines = collect_code("backend/src", [".ts"])

# Generate doc
output_path = r"d:\GitHub-lslife-V6.0\清远同城_源代码鉴别材料_60页.docx"
create_docx(android_lines, backend_lines, output_path)
print("SUCCESS")
