import os
import docx
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
import re

def remove_control_characters(s):
    return ''.join(c for c in s if (ord(c) >= 32 and ord(c) != 127) or c in '\r\n\t')

def add_page_number(run):
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')
    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = 'PAGE'
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'end')
    run._r.append(fldChar1)
    run._r.append(instrText)
    run._r.append(fldChar2)

def create_professional_manual(output_file):
    doc = docx.Document()
    
    # Base Style
    style = doc.styles['Normal']
    font = style.font
    font.name = '宋体'
    font.size = Pt(12) # 小四
    rFonts = font.element.rPr.rFonts
    if rFonts is not None:
        rFonts.set(qn('w:eastAsia'), '宋体')
    
    # Page Setup
    for section in doc.sections:
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin = Inches(1.2)
        section.right_margin = Inches(1.2)
        
        # Header
        header = section.header
        hp = header.paragraphs[0]
        hp.text = "清远智慧同城生活服务平台 V1.0 - 软件使用与设计说明书"
        hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        hp.style.font.size = Pt(9)
        run = hp.add_run("  第 ")
        add_page_number(run)
        hp.add_run(" 页")

    # --- Cover Page ---
    for _ in range(5):
        doc.add_paragraph()
        
    title = doc.add_paragraph('清远智慧同城生活服务平台\n（简称：同城清远）\n\n软件使用与设计说明书')
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in title.runs:
        run.font.size = Pt(26)
        run.bold = True
        
    for _ in range(10):
        doc.add_paragraph()
        
    info = doc.add_paragraph('版本号：V1.0\n权利人：许路\n\n')
    info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in info.runs:
        run.font.size = Pt(16)
        
    doc.add_page_break()
    
    # --- Content ---
    def add_h1(text):
        h = doc.add_heading(text, level=1)
        h.alignment = WD_ALIGN_PARAGRAPH.LEFT
        for run in h.runs:
            run.font.name = '黑体'
            rFonts = run.element.rPr.rFonts
            if rFonts is not None: rFonts.set(qn('w:eastAsia'), '黑体')
            run.font.color.rgb = RGBColor(0, 0, 0)
            
    def add_h2(text):
        h = doc.add_heading(text, level=2)
        for run in h.runs:
            run.font.name = '黑体'
            rFonts = run.element.rPr.rFonts
            if rFonts is not None: rFonts.set(qn('w:eastAsia'), '黑体')
            run.font.color.rgb = RGBColor(0, 0, 0)
            
    def add_p(text):
        p = doc.add_paragraph(text)
        p.paragraph_format.line_spacing = Pt(20) # Fixed line spacing for >= 30 lines/page
        p.paragraph_format.space_after = Pt(6)
        p.paragraph_format.first_line_indent = Pt(24) # 2 chars indent
        return p
        
    def add_img(img_path, caption):
        if os.path.exists(img_path):
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = p.add_run()
            r.add_picture(img_path, width=Inches(3.2))
            
            cp = doc.add_paragraph(caption)
            cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
            cp.paragraph_format.space_after = Pt(12)
            for r in cp.runs:
                r.font.size = Pt(10)
                r.font.color.rgb = RGBColor(100, 100, 100)

    # 1. Overview
    add_h1("一、 软件概述")
    add_p("清远智慧同城生活服务平台（简称：同城清远）是一款定位于县域下沉市场的纯净版同城分类信息与生活服务撮合平台。")
    add_p("本软件旨在为广大用户提供一个便捷、高效、安全的本地生活服务信息发布与检索渠道。涵盖二手闲置、房屋租售、求职招聘、家政服务、便民维修等多个垂直领域。通过严格的实名认证与风控审核机制，保障平台信息的真实性与可靠性。")
    
    # 2. Environment
    add_h1("二、 运行环境")
    add_h2("2.1 硬件环境")
    add_p("服务端开发环境：CPU Intel i5及以上；内存 16GB及以上；硬盘 1TB及以上。")
    add_p("服务端运行环境：4核CPU，8GB内存及以上。")
    add_p("客户端运行环境：ARM架构智能手机，4GB内存及以上。")
    add_h2("2.2 软件环境")
    add_p("开发操作系统：Windows 10/11 64位。")
    add_p("服务端操作系统：Linux (Ubuntu/CentOS)。")
    add_p("客户端操作系统：Android 7.0 及以上版本。")
    add_p("支撑环境：PostgreSQL 14 数据库，Node.js 20 运行环境。")
    
    # 3. User Guide
    add_h1("三、 软件功能说明与操作手册")
    
    add_h2("3.1 首页展示与信息检索")
    add_p("打开软件后，默认进入系统主界面。首页集成了各类生活服务的快速入口（如家政、维修、租房等），并根据用户地理位置和偏好，智能推荐附近的优质同城信息。")
    add_img("docs/images/media_1788307599284.jpg", "图1：清远智慧同城生活服务平台 - 系统主界面，包含分类服务入口及附近信息流推荐")
    
    add_h2("3.2 详细分类浏览")
    add_p("用户可以通过底部导航栏进入“分类”页面。分类页面采用层级结构设计，左侧为一级分类，右侧展示二级和三级细分服务模块。")
    add_img("docs/images/media_1788307601243.jpg", "图2：清远智慧同城生活服务平台 - 细分类目大全页面，包含家政、便民等本地服务模块")
    add_p("点击具体的三级类目（如日常保洁），系统会进入该分类下的详细信息列表，支持根据最新发布、距离远近等条件进行过滤筛选。")
    add_img("docs/images/media_1788307602888.jpg", "图3：清远智慧同城生活服务平台 - 具体服务信息的过滤展示列表")
    
    add_h2("3.3 信息发布功能")
    add_p("点击底部导航栏中间的“+”发布按钮，系统会弹出首层发布分类选择菜单。用户需明确所发布信息的所属大类。")
    add_img("docs/images/media_1788307607147.jpg", "图4：清远智慧同城生活服务平台 - 信息发布时的首层分类选择菜单")
    add_p("随后进入二级和三级具体细分类目的选择。这一设计有助于系统后续渲染专属的动态表单结构。")
    add_img("docs/images/media_1788308430991.jpg", "图5：清远智慧同城生活服务平台 - 信息发布具体细分类目选择")
    add_p("进入发布表单页面后，用户可填写标题、详细描述、价格等信息，并上传多张实拍图片。系统底层集成了 AI 智能文案润色功能，帮助用户提高信息的吸引力。")
    add_img("docs/images/media_1788307609828.jpg", "图6：清远智慧同城生活服务平台 - 信息发布表单界面与多图上传")
    
    add_h2("3.4 消息互动与即时通讯")
    add_p("在信息详情页，买家可以通过点击“联系”按钮，与信息发布者建立即时通讯会话。所有的历史聊天记录、系统通知等均汇总于底部的“消息”模块中。")
    add_img("docs/images/media_1788308433342.jpg", "图7：清远智慧同城生活服务平台 - 消息中心与即时通讯会话列表")
    
    add_h2("3.5 个人中心与系统设置")
    add_p("进入底部的“我的”页面，用户可以查看自己的个人资料、粉丝数、关注数，并管理自己发布的信息。同时包含商业推广与账户余额等高级功能。")
    add_img("docs/images/media_1788308435663.jpg", "图8：清远智慧同城生活服务平台 - 个人中心页面及管理入口")
    add_p("点击右上角的设置图标，可进入“设置与隐私”页面，进行手机号绑定、密码修改、消息通知提醒方式配置以及缓存清理等操作。")
    add_img("docs/images/media_1788308439366.jpg", "图9：清远智慧同城生活服务平台 - 账号安全、消息与隐私保护配置界面")

    # 4. Architecture
    add_h1("四、 系统架构设计")
    add_p("本项目采用前后端分离的技术架构，保证了系统的高内聚低耦合。")
    
    md_files = [
        "docs/FULL_STACK_ANALYSIS_2026-07-20.md",
        "README_V2.2_HANDOVER_WHITEPAPER.md"
    ]
    for md in md_files:
        if os.path.exists(md):
            with open(md, 'r', encoding='utf-8') as f:
                for line in f:
                    line = line.strip()
                    if not line or line.startswith('//') or line.startswith('![') or line.startswith('```') or line in ['end', 'subgraph']:
                        continue
                    line = re.sub(r'[*#>`\-]', '', line).strip()
                    if line:
                        add_p(line)
                        
    # 5. Database
    add_h1("五、 数据库结构设计")
    schema_path = "backend/prisma/schema.prisma"
    if os.path.exists(schema_path):
        add_p("本系统采用 PostgreSQL 关系型数据库，利用 Prisma ORM 进行数据建模。以下为核心业务实体与表结构设计说明：")
        with open(schema_path, 'r', encoding='utf-8') as f:
            for line in f:
                line = remove_control_characters(line.strip())
                if line and not line.startswith('//'):
                    add_p(line)
                    
    doc.save(output_file)

output_path = r"D:\GitHub-lslife-V6.0\清远同城_文档鉴别材料_V4_专业排版.docx"
create_professional_manual(output_path)
print("SUCCESS")
