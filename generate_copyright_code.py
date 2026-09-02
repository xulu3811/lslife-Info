import os
import docx
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

def collect_code():
    lines = []
    # Directories to prioritize: android front-end first, then backend
    targets = [
        "android/app/src/main/java/com/lianshan/lslife/feature",
        "backend/src"
    ]
    
    for target in targets:
        for root, dirs, files in os.walk(target):
            for file in files:
                if file.endswith('.kt') or file.endswith('.ts') or file.endswith('.tsx'):
                    filepath = os.path.join(root, file)
                    try:
                        with open(filepath, 'r', encoding='utf-8') as f:
                            for line in f:
                                line = line.rstrip()
                                # Skip empty lines and simple comment lines
                                stripped = line.strip()
                                if not stripped or stripped.startswith('//') or stripped.startswith('/*') or stripped.startswith('*'):
                                    continue
                                # Hard wrap very long lines to avoid them wrapping in Word and breaking the 50-line count physically
                                if len(line) > 80:
                                    line = line[:80] + "..."
                                lines.append(line)
                    except Exception:
                        pass
    return lines

def create_docx(code_lines, output_file):
    doc = docx.Document()
    
    # Set page margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(0.8)
        section.bottom_margin = Inches(0.8)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)
        
        # Add header
        header = section.header
        header_para = header.paragraphs[0]
        header_para.text = "清远智慧同城生活服务平台项目 V1.0"
        header_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        
        # Add page numbers to header
        run = header_para.add_run("  第 ")
        fldChar1 = OxmlElement('w:fldChar')
        fldChar1.set(qn('w:fldCharType'), 'begin')
        instrText = OxmlElement('w:instrText')
        instrText.set(qn('xml:space'), 'preserve')
        instrText.text = 'PAGE'
        fldChar2 = OxmlElement('w:fldChar')
        fldChar2.set(qn('w:fldCharType'), 'end')
        
        run._r.append(fldChar1)
        run._r.append(instrText)
        run._r.append(fldChar2)
        run = header_para.add_run(" 页")

    if len(code_lines) < 3000:
        print(f"Not enough code lines! Found only {len(code_lines)}")
        return
        
    first_1500 = code_lines[:1500]
    last_1500 = code_lines[-1500:]
    final_lines = first_1500 + last_1500
    
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Courier New'
    font.size = Pt(9)
    rFonts = font.element.rPr.rFonts
    if rFonts is not None:
        rFonts.set(qn('w:eastAsia'), '宋体')

    for page in range(60):
        page_lines = final_lines[page*50 : (page+1)*50]
        
        p = doc.add_paragraph()
        p.paragraph_format.line_spacing = Pt(12) # Fixed line spacing to fit 50 lines
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.space_before = Pt(0)
        
        p.add_run('\n'.join(page_lines))
        
        if page < 59:
            doc.add_page_break()
            
    doc.save(output_file)
    print(f"Generated {output_file} successfully.")

print("Collecting code lines...")
lines = collect_code()
print(f"Collected {len(lines)} lines.")
output_path = r"D:\LsLife\清远同城_源代码鉴别材料.docx"
create_docx(lines, output_path)
